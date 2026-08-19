from __future__ import annotations

import os
import re
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_LOGO_PATH = BASE_DIR / "logo.png"
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-mini")


@dataclass(frozen=True)
class ParticipationDocuments:
    terms_path: Path
    caption_path: Path
    question: str
    caption: str


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^\w\- ]+", "", value.strip(), flags=re.UNICODE)
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned or "Teilnahmebedingungen"


def parse_game_day(value: str | date) -> date:
    if isinstance(value, date):
        return value
    value = value.strip()
    for date_format in ("%Y-%m-%d", "%d.%m.%Y"):
        try:
            return datetime.strptime(value, date_format).date()
        except ValueError:
            continue
    raise ValueError("Bitte gib das Spieldatum im Format TT.MM.JJJJ oder JJJJ-MM-TT an.")


def get_openai_client():
    from openai import OpenAI

    return OpenAI()


def create_question(club_name: str, game_day: date, client=None) -> str:
    client = client or get_openai_client()
    response = client.responses.create(
        model=OPENAI_MODEL,
        input=(
            "Erstelle eine kurze, kreative und ansprechende Frage fuer ein Social Media "
            f"Gewinnspiel fuer das Heimspiel des {club_name} am {game_day:%d.%m.%Y}. "
            "Die Frage soll leicht sein, gut loesbar fuer Fans, maximal einen Satz lang sein, "
            "spannend und unterhaltsam formuliert sein sowie Fans aktivieren und zur Teilnahme "
            "motivieren. Sie soll kein langweiliges Trivia sein, sondern einen Bezug zu Emotion, "
            "Fanwissen oder bekannten Vereinsfakten haben. Wichtig ist, dass die Frage eindeutig "
            "formuliert ist und keine Mehrdeutigkeit zulaesst. Es darf nur genau eine objektiv "
            "richtige Antwort geben. Vermeide Fragen zur Gruendung und verwende keine Gedankenstriche. "
            "Gib ausschliesslich die Frage aus, ohne zusaetzliche Erklaerungen."
        ),
    )
    return response.output_text.strip()


def create_caption(club_name: str, game_day: date, client=None) -> str:
    client = client or get_openai_client()
    response = client.responses.create(
        model=OPENAI_MODEL,
        input=(
            "Erstelle eine kurze, kreative und ansprechende Caption fuer einen Social Media Post "
            f"zum Heimspiel des {club_name} am {game_day:%d.%m.%Y}. Die Caption soll emotional, "
            "aktivierend und fan-nah formuliert sein, sodass sie die Vorfreude auf das Spiel "
            "steigert und zur Interaktion motiviert. Sie soll maximal 2 bis 3 kurze Saetze lang "
            "sein und eine klare Call-to-Action enthalten. Gib ausschliesslich die Caption aus, "
            "ohne zusaetzliche Erklaerungen."
        ),
    )
    return response.output_text.strip()


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(7)
    run.font.name = "Calibri"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_bold_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(7)
    run.bold = True
    run.font.name = "Calibri"
    paragraph.paragraph_format.space_after = Pt(6)


def configure_terms_document(doc: Document, club_name: str, logo_path: Path) -> None:
    section = doc.sections[0]
    header = section.header

    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    sect_pr = section._sectPr
    columns = sect_pr.xpath("./w:cols")
    columns = columns[0] if columns else OxmlElement("w:cols")
    if columns.getparent() is None:
        sect_pr.append(columns)
    columns.set(qn("w:num"), "2")

    table = header.add_table(rows=1, cols=2, width=Cm(19))
    table.autofit = False
    cell_left = table.cell(0, 0)
    cell_left.width = Cm(15)
    run_left = cell_left.paragraphs[0].add_run(f'Teilnahmebedingungen fuer das Gewinnspiel "Heimspieltickets {club_name}"')
    run_left.font.size = Pt(10)
    run_left.bold = True
    run_left.font.name = "Calibri"

    cell_right = table.cell(0, 1)
    cell_right.width = Cm(5)
    paragraph_right = cell_right.paragraphs[0]
    paragraph_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logo_path.exists():
        paragraph_right.add_run().add_picture(str(logo_path), width=Cm(2))


def build_terms_document(
    club_name: str,
    opponent: str,
    game_day: date,
    question: str,
    output_path: Path,
    logo_path: Path = DEFAULT_LOGO_PATH,
) -> Path:
    today = date.today()
    contest_start = game_day - timedelta(days=5)
    contest_end = contest_start + timedelta(days=2)
    claim_deadline = contest_end + timedelta(days=1)

    doc = Document()
    configure_terms_document(doc, club_name, logo_path)

    add_bold_paragraph(
        doc,
        'Diese Teilnahmebedingungen gelten fuer das Gewinnspiel "Heimspieltickets" ueber Instagram und Facebook.',
    )
    add_body_paragraph(
        doc,
        "Das Gewinnspiel wird ausschliesslich gemeinsam von der Saechsischen Lotto-GmbH, Oststrasse 105, "
        "04299 Leipzig, und der Die Sportwerk GmbH, Stechgrundstrasse 2a, 01324 Dresden, veranstaltet "
        "und organisiert und steht in keinerlei Verbindung zu Instagram und Facebook.",
    )
    add_body_paragraph(doc, f"Das Gewinnspiel findet vom {contest_start:%d.%m.%Y} bis {contest_end:%d.%m.%Y} statt.")
    add_bold_paragraph(doc, "Teilnahme")
    add_body_paragraph(
        doc,
        "An der Verlosung nehmen alle Teilnehmer ueber Instagram und Facebook teil, die auf der "
        "Sachsenlotto-Instagram-Seite https://www.instagram.com/sachsenlotto und der "
        "Sachsenlotto-Facebook-Seite https://www.facebook.com/sachsenlotto rechtzeitig bis zum "
        f"Abgabeschluss folgende Frage in den Kommentaren beantworten: {question}",
    )
    add_body_paragraph(
        doc,
        "Jeder Teilnehmer ist an der Verlosung nur mit einem Kommentar teilnahmeberechtigt. Die Teilnahme "
        "an dem Gewinnspiel ist unabhaengig von der Teilnahme an den Lotterien und Wetten der Saechsischen "
        "Lotto-GmbH. Teilnahmeberechtigt sind nur Personen ueber 18 Jahre mit Wohnsitz in Sachsen.",
    )
    add_body_paragraph(
        doc,
        "Ausgenommen von der Teilnahme sind Mitarbeiter der Saechsischen Lotto-GmbH, Mitarbeiter der Die "
        "Sportwerk GmbH und Mitarbeiter der Unternehmen, die an der Abwicklung des Gewinnspiels beteiligt sind.",
    )
    add_body_paragraph(
        doc,
        "Der Teilnehmer ist damit einverstanden, dass die Saechsische Lotto-GmbH und Die Sportwerk GmbH seinen "
        "Namen und seine E-Mail-Adresse zur Durchfuehrung des Gewinnspiels verwenden.",
    )
    add_body_paragraph(
        doc,
        "Jeder Teilnehmer ist fuer die von ihm gelieferten Inhalte verantwortlich. Kommentare duerfen nicht "
        "gegen die guten Sitten, geltendes Recht oder Rechte Dritter verstossen.",
    )
    add_body_paragraph(
        doc,
        "Verstoesst der Teilnehmer gegen diese Bestimmungen, hat er die Saechsische Lotto-GmbH von der Haftung "
        "freizustellen, einschliesslich angemessener Kosten der Rechtsverteidigung.",
    )
    add_body_paragraph(
        doc,
        "Es liegt im Ermessen der Saechsischen Lotto-GmbH und der Die Sportwerk GmbH, Teilnehmer ohne Angabe "
        "von Gruenden vom Gewinnspiel auszuschliessen oder Kommentare zu sperren oder zu loeschen.",
    )
    add_bold_paragraph(doc, "Verlosung")
    add_body_paragraph(
        doc,
        f"Es werden zwei Gewinne von Sachsenlotto verlost. Ein Gewinn umfasst zwei Eintrittskarten fuer das "
        f"Spiel {club_name} gegen {opponent} am {game_day:%d.%m.%Y}. Es werden keine Sitzplaetze garantiert. "
        "Eine Barauszahlung, ein Umtausch oder eine Uebertragung des Gewinns sind ausgeschlossen.",
    )
    add_body_paragraph(
        doc,
        "An der Verlosung nehmen alle Kommentare ueber Instagram und Facebook teil, die auf den Gewinnspiel-Post "
        f"bis zum {contest_end:%d.%m.%Y}, 12:00 Uhr eingestellt wurden. Die Ermittlung der Gewinner findet am "
        f"{contest_end:%d.%m.%Y} um 15:00 Uhr statt.",
    )
    add_bold_paragraph(doc, "Gewinne")
    add_body_paragraph(
        doc,
        "Die Gewinner werden nach einem Zufallsverfahren aus der Datei der Teilnehmer ermittelt.",
    )
    add_body_paragraph(
        doc,
        f"Die Gewinner werden am {contest_end:%d.%m.%Y} mittels persoenlicher Nachricht ueber Instagram und "
        "Facebook ueber ihren Gewinn und die Formalitaeten informiert.",
    )
    add_body_paragraph(
        doc,
        f"Der Gewinn verfaellt, wenn der Gewinner seine persoenlichen Daten nicht oder nicht rechtzeitig bis zum "
        f"{claim_deadline:%d.%m.%Y} 15:00 Uhr mitteilt.",
    )
    add_bold_paragraph(doc, "Haftung, Anwendbares Recht, Aenderung")
    add_body_paragraph(
        doc,
        "Es besteht keine Pruefungspflicht der Saechsischen Lotto-GmbH und der Die Sportwerk GmbH fuer die "
        "Beitraege von Instagram- und Facebook-Nutzern.",
    )
    add_body_paragraph(
        doc,
        "Die Saechsische Lotto-GmbH und die Die Sportwerk GmbH haften nur fuer Schaeden, welche vorsaetzlich "
        "oder grob fahrlaessig oder durch die Verletzung von wesentlichen Vertragspflichten verursacht wurden.",
    )
    add_body_paragraph(
        doc,
        "Die Veranstalter behalten sich das Recht vor, das Gewinnspiel nach eigenem Ermessen zu beenden, zu "
        "unterbrechen oder zu aendern, falls eingetretene Umstaende dies erforderlich machen.",
    )
    add_body_paragraph(doc, "Das Gewinnspiel unterliegt deutschem Recht. Der Rechtsweg ist ausgeschlossen.")
    add_body_paragraph(doc, f"Leipzig, den {today:%d.%m.%Y} Saechsische Lotto-GmbH")
    add_body_paragraph(doc, "Die Sportwerk GmbH")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def build_caption_document(caption: str, output_path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    paragraph = doc.add_paragraph(caption)
    paragraph.paragraph_format.space_after = Pt(6)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_participation_documents(
    club_name: str,
    opponent: str,
    game_day: str | date,
    output_dir: str | Path,
    question: str | None = None,
    caption: str | None = None,
    client=None,
) -> ParticipationDocuments:
    club_name = club_name.strip()
    opponent = opponent.strip()
    if not club_name:
        raise ValueError("Bitte gib einen Verein an.")
    if not opponent:
        raise ValueError("Bitte gib einen Gegner an.")

    parsed_game_day = parse_game_day(game_day)
    output_dir = Path(output_dir)
    base_name = safe_filename(club_name)

    question = question.strip() if question else create_question(club_name, parsed_game_day, client)
    caption = caption.strip() if caption else create_caption(club_name, parsed_game_day, client)
    if not question:
        raise ValueError("Die Gewinnspielfrage ist leer.")
    if not caption:
        raise ValueError("Die Caption ist leer.")

    terms_path = output_dir / f"Teilnahmebedingungen_{base_name}.docx"
    caption_path = output_dir / f"Caption_{base_name}.docx"
    build_terms_document(club_name, opponent, parsed_game_day, question, terms_path)
    build_caption_document(caption, caption_path)

    return ParticipationDocuments(
        terms_path=terms_path,
        caption_path=caption_path,
        question=question,
        caption=caption,
    )


def teilnahmebedingungen() -> None:
    club_name = input("Welcher Verein soll erstellt werden? ")
    opponent = input("Welcher Gegner? ")
    game_day = input("Wann ist der Gameday? ")
    result = generate_participation_documents(club_name, opponent, game_day, Path.cwd())
    print("Dokumente wurden erstellt!")
    print(result.terms_path)
    print(result.caption_path)
    print(result.question)


if __name__ == "__main__":
    teilnahmebedingungen()
